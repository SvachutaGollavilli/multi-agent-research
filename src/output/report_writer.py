# src/output/report_writer.py
# ─────────────────────────────────────────────
# Writes the pipeline's final report to a Word document (.docx)
# in the results/ folder. Handles name collisions by appending _1, _2 etc.
# Parses the markdown report and maps headings/bullets to proper Word styles.
# ─────────────────────────────────────────────

from __future__ import annotations

import logging
import os
import re
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

_RESULTS_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
    "results",
)


def _safe_filename(query: str) -> str:
    """Convert a query string into a safe filename (no special chars)."""
    safe = re.sub(r"[^\w\s-]", "", query.lower())
    safe = re.sub(r"[\s_-]+", "_", safe).strip("_")
    return safe[:60] or "report"


def _unique_path(base_name: str) -> str:
    """
    Return a path that doesn't exist yet.
    results/foo.docx exists → try results/foo_1.docx → foo_2.docx etc.
    """
    os.makedirs(_RESULTS_DIR, exist_ok=True)
    candidate = os.path.join(_RESULTS_DIR, f"{base_name}.docx")
    if not os.path.exists(candidate):
        return candidate
    counter = 1
    while True:
        candidate = os.path.join(_RESULTS_DIR, f"{base_name}_{counter}.docx")
        if not os.path.exists(candidate):
            return candidate
        counter += 1


def _add_metadata_table(doc: Document, result: dict, run_id: str) -> None:
    """Add a compact run-metadata block at the top of the document."""
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    meta_items = [
        ("Run ID",           run_id[:8] + "..." if run_id else "—"),
        ("Query",            result.get("query", "")),
        ("Sources found",    str(len(result.get("sources", [])))),
        ("Claims extracted", str(len(result.get("key_claims", [])))),
        ("Total tokens",     str(result.get("token_count", 0))),
        ("Total cost",       f"${result.get('cost_usd', 0.0):.6f}"),
        ("Generated",        datetime.now().strftime("%Y-%m-%d %H:%M")),
    ]

    for label, value in meta_items:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True

    doc.add_paragraph()  # spacer after table


def _add_inline_formatting(paragraph, text: str) -> None:
    """
    Parse inline **bold** and *italic* markers and add formatted runs.
    Falls back to a single plain run if no markers found.
    """
    tokens = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    if not tokens or tokens == [text]:
        paragraph.add_run(text)
        return

    for token in tokens:
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            paragraph.add_run(token)


def _parse_and_add_content(doc: Document, markdown: str) -> None:
    """
    Parse the markdown report and add content to the Word document.

    Handles:
      # Title      → Heading 1
      ## Section   → Heading 2
      ### Sub      → Heading 3
      - bullet     → List Bullet
      1. numbered  → List Number
      **bold**     → bold run
      plain text   → Normal paragraph
    """
    for line in markdown.splitlines():
        stripped = line.strip()

        if not stripped:
            continue

        if stripped.startswith("# ") and not stripped.startswith("## "):
            doc.add_heading(stripped[2:].strip(), level=1)

        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)

        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)

        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_formatting(p, stripped[2:].strip())

        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s*", "", stripped)
            p = doc.add_paragraph(style="List Number")
            _add_inline_formatting(p, text)

        else:
            p = doc.add_paragraph()
            _add_inline_formatting(p, stripped)


def _add_pipeline_trace(doc: Document, result: dict) -> None:
    """Append a pipeline trace table at the end of the document."""
    trace = result.get("pipeline_trace", [])
    if not trace:
        return

    doc.add_heading("Pipeline Trace", level=2)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    headers = ["Agent", "Duration (ms)", "Tokens", "Summary"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for step in trace:
        row = table.add_row()
        row.cells[0].text = step.get("agent", "")
        row.cells[1].text = str(step.get("duration_ms", 0))
        row.cells[2].text = str(step.get("tokens", 0))
        row.cells[3].text = step.get("summary", "")


def write_report(result: dict, run_id: str = "") -> str:
    """
    Write the final pipeline report to a Word document.

    Args:
        result:  The final LangGraph state dict
        run_id:  The pipeline run ID (for metadata block)

    Returns:
        Absolute path to the saved .docx file
    """
    query       = result.get("query", "unknown_query")
    report_md   = result.get("current_draft") or result.get("final_report", "")
    base_name   = _safe_filename(query)
    output_path = _unique_path(base_name)

    doc = Document()

    # Title
    title = doc.add_heading("Research Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(f'Query: "{query}"')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.italic = True

    doc.add_paragraph()

    # Run metadata
    _add_metadata_table(doc, result, run_id)

    # Main report body
    if report_md:
        _parse_and_add_content(doc, report_md)
    else:
        doc.add_paragraph("No report content was generated for this run.")

    # Pipeline trace on a new page
    doc.add_page_break()
    _add_pipeline_trace(doc, result)

    doc.save(output_path)
    logger.info(f"[report_writer] saved → {output_path}")
    return output_path
